import os
import sys
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# Paths and files
source_folder = "./"
destination_file = "./thesummary-1.docx"
data_copied = "copylist.txt"

# variables
copied = []

with open(data_copied, "r") as f:
    copied = f.read().splitlines()


exceptions = ["Thoughts", "./thesummary-1.docx", "copylist.txt"]


# Check if the destination file exists and is valid
try:
    if os.path.exists(destination_file):
        print(f"\nOpening existing destination file: {destination_file}")
        endpoint_file = Document(destination_file)  # Attempt to open it
    else:
        print(f"\nCreating new destination file: {destination_file}")
        endpoint_file = Document()  # Create a new document
        endpoint_file.save(destination_file)  # Save it for later use
except PackageNotFoundError:
    print(f"\nError: {destination_file} is not a valid .docx file.")
    print("\nRecreating the file...")
    endpoint_file = Document()  # Create a new document
    endpoint_file.save(destination_file)  # Overwrite with a new valid file


def copy_data(file_path, exceptions, copied):  # Copy data from individual .docx files
    try:
        if file_path.endswith(".docx") and os.path.isfile(file_path):
            if file_path not in copied and file_path not in exceptions:
                print(f"Copying data from {file_path}")
                source_file = Document(file_path)
                endpoint_file.add_heading(f"Contets of: {file_path}", 0)
                for para in source_file.paragraphs:
                    endpoint_file.add_paragraph(para.text)
                endpoint_file.save(destination_file)
                copied.append(file_path)
    except Exception as e:
        print(f"Error copying data from {file_path}: {e}")


def process_folder(folder_path, exceptions, copied):  # Process folders recursively
    if os.path.isdir(folder_path):
        print(f"Entering folder: {folder_path}")
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)
            if os.path.isdir(item_path) and item not in exceptions:
                # Recurse into subdirectory
                process_folder(item_path, exceptions, copied)
            elif os.path.isfile(item_path) and item_path.endswith(".docx"):
                # Copy data if it's a .docx file
                copy_data(item_path, exceptions, copied)
    else:
        print(f"{folder_path} is not a valid directory.")


def main():
    process_folder(source_folder, exceptions, copied)

    with open(data_copied, "w") as f:
        for item in copied:
            f.write(item + "\n")

    print(f"\nData copied to {destination_file}")


if __name__ == "__main__":
    main()
